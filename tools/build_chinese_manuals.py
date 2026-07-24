from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
MAIN_NAME = "CVTA公文OCR自動化歸納系統_使用手冊.docx"
GOOGLE_NAME = "Google授權申請與網路設定指南.docx"

FONT = "Microsoft JhengHei"
NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "9C6500"
RED = "9C0006"
INK = "202124"
MUTED = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    paragraph.add_run(" 頁")


def add_hyperlink(paragraph, text, target, external=True):
    part = paragraph.part
    rid = part.relate_to(target, RELATIONSHIP_TYPE.HYPERLINK, is_external=external)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def apply_styles(doc, short_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 8),
        ("Subtitle", 12, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 8),
        ("Heading 2", 13, BLUE, 13, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)

    header = section.header.paragraphs[0]
    header.text = short_title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc, kicker, title, subtitle, version="版本 3.0｜2026 年 7 月"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker)
    r.bold = True
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(GOLD)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(68)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(version)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("適用於 Windows GUI 與開發環境 CLI")
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_callout(doc, label, text, kind="info"):
    colors = {
        "info": (LIGHT_BLUE, NAVY),
        "warning": ("FFF2CC", GOLD),
        "danger": ("FCE8E6", RED),
    }
    fill, color = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.6)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}　")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        widths = [6.6 / len(headers)] * len(headers)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_steps(doc, steps):
    for title, detail in steps:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(title)
        r.bold = True
        if detail:
            p.add_run(f"：{detail}")


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_contents(doc, entries):
    p = doc.add_paragraph("文件導覽", style="Heading 1")
    add_bookmark(p, "contents", 1)
    for title, description in entries:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(title)
        r.bold = True
        p.add_run(f"｜{description}")


def build_main():
    doc = Document()
    apply_styles(doc, "CVTA 公文 OCR 自動化歸納系統｜使用手冊")
    add_cover(
        doc,
        "使用者操作手冊",
        "CVTA 公文 OCR 自動化歸納系統",
        "從影像匯入、AI 分析、人工校正，到 Google Workspace 同步",
    )
    add_contents(doc, [
        ("1. 系統用途與處理流程", "了解輸入、分析、校正、保存與雲端同步。"),
        ("2. 使用前準備", "確認程式、Ollama 與輸入檔案。"),
        ("3. 初次設定", "設定 AI 服務；Google 設定另見獨立指南。"),
        ("4. GUI 日常操作", "依序完成匯入、分析、校正與上傳。"),
        ("5. 欄位與驗證規則", "掌握日期、字別、文號與關鍵事項規範。"),
        ("6. 資料保存與復原", "使用 JSON 封存與重開工作。"),
        ("7. CLI 與維運操作", "供開發、測試與批次執行。"),
        ("8. 常見問題", "快速排除 AI、檔名、Google 與資料問題。"),
    ])

    doc.add_heading("1. 系統用途與處理流程", 1)
    doc.add_paragraph(
        "本系統用於批次整理公文掃描影像。它會呼叫 Ollama 多模態模型擷取公文欄位，"
        "依內建規則驗證資料，讓使用者在畫面中比對原圖並修正，再將資料保存在本機，"
        "或選擇上傳影像至 Google Drive 並把欄位寫入 Google Sheets。"
    )
    add_table(doc, ["階段", "使用者動作", "系統結果"], [
        ("匯入", "選取 JPG、JPEG 或 PNG", "建立待處理清單並顯示影像"),
        ("分析", "按下「分析」", "OCR 擷取七個欄位並執行驗證"),
        ("校正", "逐筆比對原圖、修改欄位", "修改立即回寫目前資料"),
        ("保存", "按下「儲存」", "輸出 JSON 封存檔，供下次重開"),
        ("同步", "按下「上傳」", "影像上傳 Drive，資料追加至 Sheets"),
    ], [1.0, 2.35, 3.25])
    add_callout(
        doc, "重要",
        "Google Workspace 留空或尚未授權時，不影響本機匯入、AI 分析、人工校正與 JSON 封存。",
        "info",
    )

    doc.add_heading("2. 使用前準備", 1)
    doc.add_heading("2.1 一般使用者", 2)
    add_bullets(doc, [
        "Windows 電腦，並取得完整的程式發行資料夾；請勿只複製單一 EXE。",
        "可連線的 Ollama 服務，且已安裝設定畫面指定的模型；預設模型為 qwen3.5:9b。",
        "待處理影像格式為 .jpg、.jpeg 或 .png。",
        "檔名必須以六碼數字及底線開頭，例如 115060_0001.jpg。",
    ])
    doc.add_heading("2.2 開發環境", 2)
    add_bullets(doc, [
        "Python 3.13 以上。",
        "uv 套件管理工具。",
        "專案根目錄可讀寫，供 settings.cfg、logs 與 data 輸出使用。",
    ])
    add_callout(
        doc, "檔名規則",
        "系統只接受「六碼數字_其餘名稱.副檔名」。流水號取前六碼的第 4 至第 6 碼；"
        "不符合規則或已匯入的完整檔案路徑會被略過。",
        "warning",
    )

    doc.add_heading("3. 初次設定", 1)
    doc.add_heading("3.1 開啟設定視窗", 2)
    doc.add_paragraph(
        "首次建立設定檔時，程式會自動開啟「系統設定」。日後可從主畫面的齒輪按鈕再次開啟。"
        "儲存後，新設定會寫入 src/config/settings.cfg；重新建立 AI 後端或重新啟動程式後套用。"
    )
    add_table(doc, ["欄位", "建議值或用途"], [
        ("Ollama 主機", "本機通常為 http://127.0.0.1:11434；遠端服務請填完整 URL"),
        ("模型名稱", "必須與 Ollama 已安裝模型名稱完全一致"),
        ("可用性確認逾時", "預設 5 秒；跨網路可酌量提高"),
        ("分析請求逾時", "預設 300 秒；大型影像或慢速主機可提高"),
        ("提示詞設定檔", "指向有效的 prompts.json"),
        ("試算表 ID", "Google Sheets 網址 /d/ 與 /edit 之間的字串"),
        ("工作表名稱", "目標分頁名稱；應與實際 Sheets 分頁一致"),
        ("Drive 目標資料夾", "要查找並上傳公文影像的資料夾名稱"),
    ], [1.75, 4.85])
    doc.add_heading("3.2 Google Workspace 授權與網路設定", 2)
    p = doc.add_paragraph()
    p.add_run("本章僅保留入口。完整申請步驟、API 啟用、OAuth 同意畫面、服務帳戶、檔案分享、"
              "首次瀏覽器授權、防火牆與錯誤排除，請開啟：").bold = False
    add_hyperlink(p, "Google授權申請與網路設定指南.docx", GOOGLE_NAME)
    add_callout(
        doc, "安全提醒",
        "google_key.json、client_secret.json 與 token.json 均屬敏感憑證，不得寄送、上傳版本控制或放入公開共享空間。",
        "danger",
    )

    doc.add_heading("4. GUI 日常操作", 1)
    doc.add_heading("4.1 啟動程式", 2)
    add_steps(doc, [
        ("發行版", "開啟 CVTA-Documents-Analyzer 資料夾，再執行 CVTA-Documents-Analyzer.exe。"),
        ("開發版", "在專案根目錄執行 uv run python -m src.luncher。"),
        ("確認設定", "若自動顯示設定視窗，先完成 Ollama 主機、模型與提示詞路徑。"),
    ])
    doc.add_heading("4.2 匯入公文影像", 2)
    add_steps(doc, [
        ("按下新增按鈕", "在檔案選擇器中一次選取一張或多張影像。"),
        ("等待匯入", "進度視窗會顯示目前檔名；可按「取消」停止後續檔案。"),
        ("檢查結果", "若有檔名錯誤或重複檔案，系統會列出未匯入項目。"),
        ("選取資料列", "左側顯示原始影像，右側顯示可編輯欄位。"),
    ])
    doc.add_heading("4.3 執行 AI 分析", 2)
    add_steps(doc, [
        ("先確認 Ollama", "服務須可連線，且指定模型已安裝。"),
        ("按下分析按鈕", "系統批次處理資料庫內尚待分析的公文。"),
        ("等待完成摘要", "畫面顯示總筆數、成功筆數與失敗筆數。"),
        ("逐筆複核", "AI 結果可能誤判；必須以原圖為準人工核對。"),
    ])
    add_callout(
        doc, "日期限制",
        "公文日期須為 YYYY-MM-DD，且與執行當日相差不得超過設定的 90 天；超出時該筆驗證不通過。",
        "warning",
    )
    doc.add_heading("4.4 人工校正與影像檢視", 2)
    add_bullets(doc, [
        "選取表格資料列後，可修改發文單位、發文日期、承辦人、發文字、發文號、班級與事由。",
        "修改欄位後會更新目前公文資料；切換資料列前仍建議確認內容。",
        "影像可用 Ctrl + 滑鼠滾輪縮放，也可拖曳檢視放大區域。",
        "上一筆／下一筆按鈕可循環切換資料列。",
    ])
    doc.add_heading("4.5 上傳 Google Workspace", 2)
    add_steps(doc, [
        ("先完成授權", "依獨立指南放置憑證、分享 Drive 資料夾與 Sheets。"),
        ("確認資料", "建議先完成分析與人工校正。"),
        ("按下上傳", "系統先處理 OAuth token，再上傳影像並追加試算表資料。"),
        ("處理同名檔", "依對話框選擇略過、更新既有檔案或另行上傳。"),
        ("檢查摘要", "確認成功、失敗與取消筆數，並抽查 Drive 連結及 Sheets 新增列。"),
    ])

    doc.add_heading("5. 欄位與驗證規則", 1)
    add_table(doc, ["畫面欄位", "資料鍵", "規則"], [
        ("發文日期", "doc_date", "YYYY-MM-DD；不得空白；日期差須在 90 天內"),
        ("發文字", "doc_category", "字別，例如「中分署訓」；不得空白"),
        ("發文號", "doc_number", "文號純數字；不得空白"),
        ("發文單位", "doc_from", "機關全銜；不得空白"),
        ("承辦人", "case_officer", "不得空白"),
        ("班級", "related_class", "不得空白"),
        ("事由", "key_points", "至少一項；陣列或逗號分隔文字，保存時轉成字串"),
    ], [1.2, 1.45, 3.95])
    doc.add_paragraph(
        "AI 分析只在所有必填欄位及日期規則通過時標記成功。若某筆失敗，可查看 logs/app.log 與 "
        "logs/error.log，並在 GUI 中依原圖補正後再保存。"
    )

    doc.add_heading("6. 資料保存與復原", 1)
    doc.add_heading("6.1 儲存封存檔", 2)
    add_steps(doc, [
        ("按下儲存按鈕", "選擇 JSON 檔案位置與名稱。"),
        ("保存工作狀態", "封存內容包括欄位、影像來源、分析狀態與 Drive 檔案 ID。"),
        ("保留原始影像", "JSON 只記錄影像路徑，不會把影像嵌入封存檔；請勿任意移動影像。"),
    ])
    doc.add_heading("6.2 開啟封存檔", 2)
    add_steps(doc, [
        ("按下開啟按鈕", "選擇先前儲存的 JSON。"),
        ("確認影像路徑", "若影像被移動或刪除，預覽與後續上傳可能失敗。"),
        ("檢查表格內容", "重新選取資料列並抽查欄位。"),
    ])
    doc.add_heading("6.3 清除目前資料", 2)
    doc.add_paragraph(
        "按下重設／清除按鈕會清空目前資料清單。若資料仍需保留，請先另存 JSON 封存檔。"
    )

    doc.add_heading("7. CLI 與維運操作", 1)
    add_table(doc, ["目的", "PowerShell 指令"], [
        ("同步依賴", "uv sync"),
        ("啟動 GUI", "uv run python -m src.luncher"),
        ("批次分析", "uv run python -m src.domain.orchestrator --images <image1> <image2>"),
        ("額外輸出", "uv run python -m src.domain.orchestrator --images <image1> --extra_output"),
        ("顯示說明", "uv run python -m src.domain.orchestrator --help"),
        ("執行測試", "uv run python -m unittest discover -s tests -v"),
    ], [1.55, 5.05])
    doc.add_heading("7.1 日誌與輸出", 2)
    add_bullets(doc, [
        "一般執行資訊與警告：logs/app.log。",
        "錯誤與例外：logs/error.log。",
        "CLI CSV 與額外輸出：依 Storage 設定寫入 data/output_results。",
        "設定真相來源：src/config/settings.py、src/config/settings.cfg 與 src/config/prompts.json。",
    ])

    doc.add_heading("8. 常見問題", 1)
    add_table(doc, ["現象", "處理方式"], [
        ("AI 服務無法連線", "確認 Ollama 已啟動、主機 URL 可達、防火牆允許連線。"),
        ("指定模型不存在", "在 Ollama 主機安裝同名模型，或修改設定中的模型名稱。"),
        ("圖片無法匯入", "確認格式及「六碼數字_」檔名前綴，並排除重複路徑。"),
        ("分析顯示日期錯誤", "改為 YYYY-MM-DD，並確認與當日相差不超過 90 天。"),
        ("Google 認證失敗", f"依 {GOOGLE_NAME} 檢查三個 JSON、API、分享權限與測試使用者。"),
        ("上傳找不到資料夾", "確認目標資料夾名稱完全一致，且服務帳戶已取得權限。"),
        ("Sheets 無新增資料", "確認 Spreadsheet ID、工作表名稱、服務帳戶編輯權與 Sheets API。"),
        ("封存後看不到影像", "把影像移回原路徑，或重新匯入正確位置的影像。"),
    ], [1.75, 4.85])
    add_callout(
        doc, "支援資訊",
        "回報問題時請附上操作時間、畫面訊息、相關 logs/app.log 與 logs/error.log 片段；請先移除憑證、token 與個資。",
        "info",
    )

    doc.core_properties.title = "CVTA 公文 OCR 自動化歸納系統使用手冊"
    doc.core_properties.subject = "中文使用者操作手冊"
    doc.core_properties.author = "CVTA Documents Analyzer 專案"
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUT / MAIN_NAME)


def build_google():
    doc = Document()
    apply_styles(doc, "Google 授權申請與網路設定指南")
    add_cover(
        doc,
        "獨立設定指南",
        "Google 授權申請與網路設定",
        "CVTA 公文 OCR 自動化歸納系統專用",
    )
    add_contents(doc, [
        ("1. 架構與申請清單", "理解服務帳戶與桌面版 OAuth 的分工。"),
        ("2. 建立 Google Cloud 專案", "建立或選取專案並啟用 API。"),
        ("3. 設定 OAuth 同意畫面", "設定受眾、應用程式與測試使用者。"),
        ("4. 建立兩類憑證", "下載服務帳戶金鑰與桌面版 OAuth 用戶端。"),
        ("5. 準備 Drive 與 Sheets", "建立資源、分享權限與取得 ID。"),
        ("6. 放置檔案並首次授權", "依固定檔名放置 JSON 並產生 token。"),
        ("7. 網路與防火牆", "確認 Google 與 Ollama 連線。"),
        ("8. 驗收與疑難排解", "依錯誤訊息逐項檢查。"),
        ("9. 官方網路文件", "集中列出申請與管理入口。"),
    ])

    doc.add_heading("1. 架構與申請清單", 1)
    doc.add_paragraph(
        "本程式不是只用一組 Google 憑證。請準備「服務帳戶 JSON」與「桌面版 OAuth 用戶端 JSON」兩類憑證；"
        "程式首次執行 OAuth 後，會另產生 token.json。"
    )
    add_table(doc, ["元件", "程式檔名", "用途與權限範圍"], [
        ("服務帳戶金鑰", "google_key.json", "Drive 與 Sheets API；服務帳戶需被直接分享目標資源"),
        ("桌面版 OAuth 用戶端", "client_secret.json", "啟動本機瀏覽器授權流程"),
        ("使用者權杖", "token.json", "由程式自動產生；使用 drive.file 範圍執行上傳"),
    ], [1.55, 1.8, 3.25])
    add_callout(
        doc, "不需要網域全權委派",
        "本系統只存取明確分享的 Drive 資料夾與 Google Sheets，不需替整個 Workspace 網域模擬使用者。"
        "Google 官方也建議此情境直接把特定檔案或資料夾分享給服務帳戶。",
        "info",
    )
    doc.add_heading("1.1 完成前準備", 2)
    add_bullets(doc, [
        "一個可管理 Google Cloud 專案的 Google 帳戶。",
        "一個將在瀏覽器完成 OAuth 授權、且可上傳 Drive 檔案的 Google 帳戶。",
        "目標 Google Drive 資料夾，以及目標 Google Sheets 試算表。",
        "若為 Workspace 組織帳戶，確認管理員未封鎖第三方 OAuth、外部分享或服務帳戶分享。",
    ])

    doc.add_heading("2. 建立 Google Cloud 專案並啟用 API", 1)
    add_steps(doc, [
        ("登入 Google Cloud Console", "前往 console.cloud.google.com。"),
        ("建立或選取專案", "建議使用組織管理的專用專案，名稱可設為 CVTA Documents Analyzer。"),
        ("開啟 API 程式庫", "在「API 和服務」中搜尋並啟用 Google Drive API。"),
        ("啟用 Google Sheets API", "同一專案內再搜尋並啟用 Google Sheets API。"),
        ("確認專案一致", "後續 OAuth 用戶端與服務帳戶都要建立在這個專案內。"),
    ])
    p = doc.add_paragraph("官方說明：")
    add_hyperlink(p, "啟用 Google Workspace API", "https://developers.google.com/workspace/guides/enable-apis")
    doc.add_paragraph(
        "API 啟用可能需要數分鐘生效。若程式回報 API 尚未使用或被停用，先回到該專案的 API 控制台確認。"
    )

    doc.add_heading("3. 設定 OAuth 同意畫面", 1)
    add_steps(doc, [
        ("開啟 Google Auth Platform", "依序設定 Branding、Audience 與 Data Access。"),
        ("填寫應用程式資訊", "至少填入應用程式名稱、使用者支援電子郵件與開發人員聯絡信箱。"),
        ("選擇受眾", "僅組織內使用可選 Internal；其他情境選 External。"),
        ("加入範圍", "本程式的使用者 OAuth 需要 https://www.googleapis.com/auth/drive.file。"),
        ("加入測試使用者", "若 External 應用仍為 Testing，將實際登入授權的帳戶加入 Test users。"),
        ("保存設定", "測試階段無須為少數內部測試者立即發布，但授權帳戶必須在允許名單。"),
    ])
    add_callout(
        doc, "測試模式注意",
        "若看到「存取權遭封鎖」或應用程式未完成驗證，先確認登入帳戶已列為測試使用者；"
        "正式提供廣泛外部使用前，再依 Google 規範評估品牌或 OAuth 驗證。",
        "warning",
    )
    p = doc.add_paragraph("官方說明：")
    add_hyperlink(p, "設定 OAuth 同意畫面", "https://developers.google.com/workspace/guides/configure-oauth-consent")
    p.add_run("；")
    add_hyperlink(p, "OAuth 正式環境準備", "https://developers.google.com/identity/protocols/oauth2/production-readiness/policy-compliance")

    doc.add_heading("4. 建立兩類憑證", 1)
    doc.add_heading("4.1 建立服務帳戶與 JSON 金鑰", 2)
    add_steps(doc, [
        ("前往 IAM 與管理 > 服務帳戶", "選取同一 Google Cloud 專案。"),
        ("建立服務帳戶", "輸入名稱與說明；一般不必授與專案層級管理角色。"),
        ("記錄電子郵件", "格式通常為 名稱@專案ID.iam.gserviceaccount.com。"),
        ("建立金鑰", "進入該服務帳戶的「金鑰」，新增金鑰並選擇 JSON。"),
        ("立即安全保存", "下載檔案只能取得一次；重新命名為 google_key.json。"),
    ])
    add_callout(
        doc, "最小權限",
        "IAM 專案角色不會自動授與 Google Drive 或 Sheets 文件權限。請在第 5 章直接把目標資源分享給服務帳戶。",
        "info",
    )
    p = doc.add_paragraph("官方說明：")
    add_hyperlink(p, "建立 Google Workspace 存取憑證", "https://developers.google.com/workspace/guides/create-credentials")
    p.add_run("；")
    add_hyperlink(p, "建立與刪除服務帳戶金鑰", "https://cloud.google.com/iam/docs/keys-create-delete")

    doc.add_heading("4.2 建立桌面版 OAuth 用戶端", 2)
    add_steps(doc, [
        ("前往 Google Auth Platform > Clients", "確認目前仍是同一 Cloud 專案。"),
        ("建立用戶端", "Application type 選擇 Desktop app。"),
        ("命名並建立", "例如 CVTA Documents Analyzer Desktop。"),
        ("下載 JSON", "下載 OAuth 用戶端設定並重新命名為 client_secret.json。"),
    ])
    doc.add_paragraph(
        "桌面版流程由程式啟動本機暫時回呼伺服器並自動選擇可用連接埠，因此不需手動填固定的 localhost 連接埠。"
    )

    doc.add_heading("5. 準備 Drive 與 Sheets 資源", 1)
    doc.add_heading("5.1 建立與分享 Drive 資料夾", 2)
    add_steps(doc, [
        ("建立目標資料夾", "例如「公文掃描」；避免同一可見範圍內出現多個同名資料夾。"),
        ("開啟共用", "把服務帳戶電子郵件加入資料夾。"),
        ("設定編輯權", "需允許新增與管理檔案；服務帳戶沒有信箱，可取消通知。"),
        ("記錄資料夾名稱", "在系統設定的「Drive 目標資料夾」填入完全相同名稱。"),
    ])
    p = doc.add_paragraph("官方說明：")
    add_hyperlink(p, "直接把 Workspace 資源分享給服務帳戶", "https://developers.google.com/workspace/guides/create-credentials")
    p.add_run("；")
    add_hyperlink(p, "Google Drive 資料夾共用", "https://support.google.com/drive/answer/7166529")
    doc.add_heading("5.2 建立與分享 Google Sheets", 2)
    add_steps(doc, [
        ("建立試算表", "準備既有欄名與範本列；程式會追加資料並嘗試複製上一列格式。"),
        ("把試算表分享給服務帳戶", "權限設為編輯者。"),
        ("取得 Spreadsheet ID", "從網址 https://docs.google.com/spreadsheets/d/【ID】/edit 複製 ID。"),
        ("確認分頁名稱", "記錄底部分頁標籤，例如 Sheet1 或實際業務名稱。"),
        ("填入系統設定", "設定試算表 ID、工作表名稱與 Drive 目標資料夾。"),
    ])

    doc.add_heading("6. 放置憑證並完成首次授權", 1)
    doc.add_heading("6.1 固定路徑與檔名", 2)
    add_table(doc, ["檔案", "開發環境位置", "處理方式"], [
        ("google_key.json", "src/config/google_key.json", "下載服務帳戶 JSON 後改名"),
        ("client_secret.json", "src/config/client_secret.json", "下載桌面版 OAuth JSON 後改名"),
        ("token.json", "src/config/token.json", "勿預先建立；首次授權後由程式產生"),
    ], [1.65, 2.65, 2.3])
    add_callout(
        doc, "發行版",
        "打包版的實際設定目錄須與程式發行方式一致。若發行資料夾內未提供可寫入的設定位置，"
        "請由系統維護者確認 PyInstaller 規格與執行時設定路徑後再部署憑證。",
        "warning",
    )
    doc.add_heading("6.2 首次 OAuth 授權", 2)
    add_steps(doc, [
        ("啟動程式", "確保 client_secret.json 已放好且電腦可開啟預設瀏覽器。"),
        ("執行上傳", "按下「上傳」後，程式若沒有有效 token 會啟動瀏覽器。"),
        ("選擇正確帳戶", "使用已列入測試使用者、且有權使用目標 Drive 的帳戶。"),
        ("檢視權限", "同意程式使用 drive.file 範圍管理由應用程式開啟或建立的檔案。"),
        ("返回程式", "瀏覽器顯示完成後可關閉頁面；程式會寫入 token.json。"),
        ("驗證更新", "後續 token 過期時通常會以 refresh token 自動更新。"),
    ])
    add_callout(
        doc, "重新授權",
        "若更換 OAuth 用戶端、帳戶或範圍，先安全備份後移除舊 token.json，再次執行上傳完成新授權。",
        "info",
    )

    doc.add_heading("7. 網路、防火牆與代理伺服器", 1)
    add_table(doc, ["連線", "目的", "建議檢查"], [
        ("https://accounts.google.com", "登入與 OAuth 同意", "允許 HTTPS 443、瀏覽器登入與重新導向"),
        ("https://oauth2.googleapis.com", "交換與更新 token", "允許 HTTPS 443，不攔截憑證內容"),
        ("https://www.googleapis.com", "Drive 與 Sheets API", "允許 HTTPS 443 與 Google API 網域"),
        ("http://127.0.0.1:隨機連接埠", "桌面版 OAuth 回呼", "允許程式在本機短暫監聽；不需對外開放"),
        ("Ollama 主機:11434", "AI 模型分析", "依實際設定允許 TCP；遠端主機需可路由"),
    ], [2.1, 2.0, 2.5])
    add_bullets(doc, [
        "公司代理伺服器若要求認證，需確認 Python／打包程式可使用代理；瀏覽器能開網頁不代表 API 程式一定可連線。",
        "SSL 解密設備可能造成憑證驗證失敗；請由網管以組織信任鏈處理，不要停用 TLS 驗證。",
        "本機 OAuth 回呼只綁定 localhost；防火牆提示時只需允許私人網路或本機需求，不應開放公網入站。",
        "若 Ollama 位於其他電腦，請確認該服務有正確監聽位址與存取控制，避免把未保護的 API 暴露到網際網路。",
    ])

    doc.add_heading("8. 驗收與疑難排解", 1)
    doc.add_heading("8.1 上線前驗收清單", 2)
    add_bullets(doc, [
        "Drive API 與 Sheets API 均在正確 Cloud 專案中啟用。",
        "服務帳戶 JSON 與桌面版 OAuth JSON 分別命名正確，且不是同一個檔案。",
        "目標 Drive 資料夾與試算表都已分享給服務帳戶並具有編輯權。",
        "External 測試模式下，實際登入帳戶已加入測試使用者。",
        "設定畫面的 Spreadsheet ID、工作表名稱與資料夾名稱完全正確。",
        "首次上傳已產生 token.json，Drive 可見影像，Sheets 可見新增列與連結。",
        "憑證檔未納入 Git、郵件附件、聊天訊息或公開備份。",
    ])
    doc.add_heading("8.2 常見錯誤", 2)
    add_table(doc, ["錯誤或現象", "可能原因", "處理方式"], [
        ("找不到服務帳戶檔案", "google_key.json 缺少或路徑錯誤", "放至設定目錄並確認檔名"),
        ("找不到 OAuth Client Secret", "client_secret.json 缺少或格式錯誤", "下載 Desktop app JSON 並改名"),
        ("access_denied／存取遭封鎖", "非測試使用者、管理員政策或拒絕同意", "加入測試使用者並請管理員檢查 API 控制"),
        ("invalid_grant", "token 被撤銷、時間偏差或用戶端已更換", "校正系統時間並重新產生 token"),
        ("API has not been used", "Drive 或 Sheets API 未啟用", "在相同 Cloud 專案啟用對應 API"),
        ("insufficient permissions", "Drive／Sheets 未分享或角色不足", "把資源直接分享給服務帳戶並設編輯權"),
        ("找不到目標資料夾", "名稱不同、同名衝突或不可見", "核對名稱並檢查服務帳戶可見權限"),
        ("試算表追加失敗", "ID、分頁名或編輯權錯誤", "核對網址 ID、分頁標籤與分享權限"),
        ("瀏覽器授權後程式無反應", "localhost 回呼被防火牆阻擋", "允許程式本機回呼，重新執行授權"),
        ("SSL／連線逾時", "代理、TLS 攔截或網路封鎖", "請網管允許 Google API 網域與正確信任鏈"),
    ], [1.65, 2.2, 2.75])

    doc.add_heading("9. Google 官方網路文件", 1)
    doc.add_paragraph("以下連結集中提供申請與後續維運依據；Google Console 介面名稱可能更新，請以官方頁面為準。")
    sources = [
        ("建立 Google Workspace 存取憑證", "https://developers.google.com/workspace/guides/create-credentials"),
        ("設定 OAuth 同意畫面", "https://developers.google.com/workspace/guides/configure-oauth-consent"),
        ("啟用 Google Workspace API", "https://developers.google.com/workspace/guides/enable-apis"),
        ("管理 Google Workspace 憑證", "https://developers.google.com/workspace/guides/manage-credentials"),
        ("OAuth 2.0 正式環境政策與準備", "https://developers.google.com/identity/protocols/oauth2/production-readiness/policy-compliance"),
        ("建立服務帳戶", "https://cloud.google.com/iam/docs/service-accounts-create"),
        ("建立與刪除服務帳戶金鑰", "https://cloud.google.com/iam/docs/keys-create-delete"),
        ("Google Drive 資料夾共用", "https://support.google.com/drive/answer/7166529"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)
        p.add_run(f"｜{url}")
    doc.add_paragraph("文件核對日期：2026 年 7 月 24 日。")

    doc.core_properties.title = "Google 授權申請與網路設定指南"
    doc.core_properties.subject = "CVTA 公文 OCR 自動化歸納系統 Google Workspace 設定"
    doc.core_properties.author = "CVTA Documents Analyzer 專案"
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUT / GOOGLE_NAME)


if __name__ == "__main__":
    build_main()
    build_google()
    print(OUT / MAIN_NAME)
    print(OUT / GOOGLE_NAME)
